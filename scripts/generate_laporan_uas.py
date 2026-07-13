from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / 'Template Laporan Akhir Projek UAS.docx'
OUTPUT = ROOT / 'Laporan Akhir Projek UAS - SIRS.docx'

TITLE = 'SIRS - Sistem Informasi Respons Insiden Keamanan Siber'


def clear_document(doc: Document):
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith('}sectPr'):
            continue
        body.remove(child)


def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)


def set_default_font(doc):
    styles = doc.styles
    for style_name in ['Normal', 'Body Text']:
        if style_name in styles:
            style = styles[style_name]
            style.font.name = 'Times New Roman'
            style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            style.font.size = Pt(12)
    for style_name in ['Heading 1', 'Heading 2', 'Heading 3']:
        if style_name in styles:
            style = styles[style_name]
            style.font.name = 'Times New Roman'
            style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            style.font.bold = True


def fmt_paragraph(p, justify=True, space_after=6, first_line=False):
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_line:
        p.paragraph_format.first_line_indent = Inches(0.35)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(12)


def add_p(doc, text='', style=None, align=None, bold=False, italic=False, first_line=False):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
    fmt_paragraph(p, justify=(align is None), first_line=first_line)
    if align:
        p.alignment = align
    return p


def add_heading(doc, text, level=1):
    if level == 1:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(12)
        fmt_paragraph(p, justify=False, space_after=6)
    else:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(12)
        fmt_paragraph(p, justify=False, space_after=4)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.add_run(f'• {item}')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        fmt_paragraph(p, justify=True)


def add_numbers(doc, items):
    for idx, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.add_run(f'{idx}. {item}')
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        fmt_paragraph(p, justify=True)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = 'Table Grid'
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        shade_cell(hdr[i], 'D9EAF7')
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
            fmt_paragraph(p, justify=False, space_after=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cells[i].paragraphs:
                fmt_paragraph(p, justify=True, space_after=0)
    doc.add_paragraph()
    return table


def code_block(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    r.font.size = Pt(9)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(6)
    return p


def cover(doc):
    add_p(doc, 'LAPORAN UJIAN AKHIR SEMESTER\nPRAKTIKUM PEMROGRAMAN PERANGKAT BERGERAK', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    doc.add_paragraph('\n')
    add_p(doc, f'“{TITLE}”', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    doc.add_paragraph('\n\n')
    add_p(doc, 'Disusun Oleh:', align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, 'Nama Mahasiswa: ........................................', align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, 'NPM/Kelas: ........................................', align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph('\n\n')
    add_p(doc, 'PROGRAM STUDI DIV REKAYASA KEAMANAN SIBER', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_p(doc, 'JURUSAN KOMPUTER DAN BISNIS', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_p(doc, 'POLITEKNIK NEGERI CILACAP', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_p(doc, '2025/2026', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    doc.add_page_break()


def daftar(doc):
    add_heading(doc, 'Daftar Isi', 1)
    entries = [
        ('BAB I Pendahuluan', '1'), ('1.1 Latar Belakang', '1'), ('1.2 Rumusan Masalah', '2'), ('1.3 Tujuan', '2'), ('1.4 Manfaat', '3'), ('1.5 Batasan Masalah', '3'),
        ('BAB II Analisis dan Perancangan Sistem', '4'), ('2.1 Analisis Kebutuhan Sistem', '4'), ('2.2 Rancangan Arsitektur Sistem', '7'), ('2.3 Rancangan Database', '9'), ('2.4 Rancangan Antarmuka Pengguna', '11'),
        ('BAB III Implementasi dan Pembahasan', '13'), ('3.1 Lingkungan Pengembangan', '13'), ('3.2 Implementasi Antarmuka Pengguna', '14'), ('3.3 Implementasi Firebase dan Security Rules', '18'), ('3.4 Implementasi Cloudinary dan Cloudflare Worker', '22'), ('3.5 Fitur-Fitur Utama Aplikasi', '25'), ('3.6 Pengujian Aplikasi', '31'),
        ('BAB IV Kesimpulan dan Saran', '35'), ('Daftar Pustaka', '37'), ('Lampiran', '38')]
    for name, page in entries:
        add_p(doc, f'{name}\t{page}', align=WD_ALIGN_PARAGRAPH.LEFT)
    add_heading(doc, 'Daftar Gambar dan Daftar Tabel', 1)
    for item in ['Gambar 1. Diagram Arsitektur Sistem\t8', 'Gambar 2. DFD Level 0\t8', 'Gambar 3. DFD Level 1\t9', 'Tabel 1. Lingkungan Pengembangan\t13', 'Tabel 2. Kebutuhan Fungsional\t4', 'Tabel 3. Analisis Ancaman dan Mitigasi\t6', 'Tabel 4. Struktur Collection Firestore\t10', 'Tabel 5. Pengujian Fungsional\t31', 'Tabel 6. Pengujian Keamanan\t33']:
        add_p(doc, item, align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_page_break()


def bab1(doc):
    add_heading(doc, 'BAB I\nPENDAHULUAN', 1)
    add_heading(doc, '1.1 Latar Belakang', 2)
    for t in [
        'Perkembangan aplikasi mobile Android membuat proses pelaporan layanan dan insiden semakin mudah dilakukan, tetapi pada saat yang sama memperbesar risiko keamanan siber. Data laporan insiden dapat berisi identitas pengguna, deskripsi kejadian, lokasi, bukti gambar, serta riwayat penanganan. Apabila aplikasi tidak dirancang dengan kontrol keamanan yang baik, data tersebut dapat bocor, diubah oleh pihak tidak berwenang, atau disalahgunakan.',
        'Berdasarkan kebutuhan tersebut, project ini mengembangkan SIRS, yaitu Sistem Informasi Respons Insiden Keamanan Siber berbasis Android. Aplikasi ini membantu pengguna melaporkan insiden keamanan siber secara terstruktur dan membantu admin memantau, mengklasifikasi, serta memberikan respons terhadap laporan yang masuk. Tema ini dipilih karena relevan dengan bidang Rekayasa Keamanan Siber dan memerlukan penerapan autentikasi, otorisasi, validasi input, validasi upload file, audit log, dan aturan akses database.',
        'Permasalahan konkret yang diselesaikan aplikasi adalah belum tersedianya media pelaporan insiden yang terpusat, memiliki status penanganan, menyimpan bukti pendukung, serta memisahkan hak akses antara pengguna biasa dan administrator. SIRS dirancang agar setiap laporan memiliki kode laporan, kategori, severity, status, SLA, attachment, response admin, dan riwayat perubahan status.',
        'Penerapan keamanan menjadi bagian penting dalam pengembangan aplikasi Android ini. Project menggunakan Firebase Authentication untuk identitas pengguna, Cloud Firestore dengan Security Rules untuk database, Firebase App Check, validasi file berbasis MIME, ekstensi, ukuran, dan magic bytes, serta Cloudflare Worker untuk operasi Cloudinary yang membutuhkan signature. Dengan demikian, aplikasi tidak hanya menjalankan fitur CRUD, tetapi juga menunjukkan praktik pengamanan pada layer aplikasi, backend service, dan database.'
    ]: add_p(doc, t, first_line=True)
    add_heading(doc, '1.2 Rumusan Masalah', 2)
    add_numbers(doc, [
        'Bagaimana membangun aplikasi Android native untuk pelaporan dan respons insiden keamanan siber menggunakan Kotlin dan Jetpack Compose?',
        'Bagaimana mengimplementasikan autentikasi pengguna dan pemisahan hak akses antara user dan admin menggunakan Firebase Authentication dan Cloud Firestore?',
        'Bagaimana merancang struktur data laporan, kategori, status history, activity log, dan device token pada Cloud Firestore?',
        'Bagaimana menerapkan validasi keamanan upload bukti laporan dan integrasi Cloudinary tanpa membocorkan API secret di aplikasi Android?',
        'Bagaimana menguji fitur utama dan kontrol keamanan aplikasi agar sesuai dengan kebutuhan project akhir?'
    ])
    add_heading(doc, '1.3 Tujuan', 2)
    add_numbers(doc, [
        'Menghasilkan aplikasi Android SIRS yang dapat digunakan untuk membuat, melihat, mengedit, menghapus, dan memantau laporan insiden keamanan siber.',
        'Menerapkan login, registrasi, reset password, role user/admin, dan navigasi berbasis peran.',
        'Menyusun data laporan pada Cloud Firestore dengan struktur collection yang mendukung kategori, status history, activity log, dan token FCM.',
        'Mengimplementasikan validasi file dan signed upload Cloudinary melalui Cloudflare Worker sehingga secret tidak berada di mobile app.',
        'Mendokumentasikan implementasi, arsitektur, database, fitur, serta hasil pengujian aplikasi secara komprehensif.'
    ])
    add_heading(doc, '1.4 Manfaat', 2)
    add_bullets(doc, [
        'Bagi pengguna akhir, aplikasi mempermudah pelaporan insiden keamanan siber beserta bukti pendukung dan status penanganannya.',
        'Bagi admin, aplikasi menyediakan dashboard, filter laporan, response, status history, manajemen kategori, dan activity log untuk memantau penanganan insiden.',
        'Bagi mahasiswa sebagai developer, project ini menjadi latihan penerapan Android native, Jetpack Compose, Firebase, clean architecture, dependency injection, dan kontrol keamanan aplikasi mobile.'
    ])
    add_heading(doc, '1.5 Batasan Masalah', 2)
    add_bullets(doc, [
        'Aplikasi berjalan pada Android minimum SDK 26 dan target SDK 36.',
        'Backend utama menggunakan Firebase Authentication dan Cloud Firestore, bukan MySQL atau SQLite lokal.',
        'Upload bukti menggunakan Cloudinary dengan format JPG, PNG, dan PDF.',
        'Aplikasi tidak menyimpan API secret Cloudinary di Android; operasi signature/delete dilindungi backend Cloudflare Worker.',
        'Fitur biometrik dan database lokal SQLCipher tidak menjadi ruang lingkup implementasi saat ini.',
        'Pembuatan admin tidak dibuka melalui registrasi publik; role admin dikelola melalui data pengguna di Firestore.'
    ])
    doc.add_page_break()


def bab2(doc):
    add_heading(doc, 'BAB II\nANALISIS DAN PERANCANGAN SISTEM', 1)
    add_heading(doc, '2.1 Analisis Kebutuhan Sistem', 2)
    add_p(doc, 'Kebutuhan sistem disusun berdasarkan fitur yang ditemukan pada project SIRS dan dokumentasi project akhir. Aplikasi memiliki dua aktor utama, yaitu pengguna dan admin.', first_line=True)
    add_table(doc, ['No', 'Kebutuhan Fungsional', 'Aktor'], [
        [1, 'Login, registrasi, reset password, dan logout.', 'User/Admin'],
        [2, 'Dashboard pengguna untuk melihat ringkasan laporan sendiri.', 'User'],
        [3, 'Dashboard admin untuk melihat statistik seluruh laporan.', 'Admin'],
        [4, 'Membuat laporan insiden berisi judul, kategori, severity, deskripsi, lokasi, tanggal, dan bukti.', 'User/Admin'],
        [5, 'Melihat daftar, detail, filter status, filter severity, edit, dan hapus laporan pending.', 'User'],
        [6, 'Melihat seluruh laporan, mencari laporan, filter status/severity/tanggal, dan membuka detail laporan.', 'Admin'],
        [7, 'Mengubah status, severity, dan response penanganan laporan.', 'Admin'],
        [8, 'Mengelola kategori laporan dan panduan mitigasi.', 'Admin'],
        [9, 'Mencatat activity log dan status history.', 'Sistem/Admin'],
        [10, 'Menghitung batas waktu penanganan berdasarkan severity.', 'Sistem']
    ])
    add_p(doc, 'Kebutuhan non-fungsional aplikasi meliputi keamanan, kompatibilitas, performa, maintainability, dan usability. Keamanan diwujudkan melalui Firebase Auth, Firestore Rules, validasi input, validasi file, App Check, serta pemisahan secret ke backend Worker. Kompatibilitas ditentukan oleh minSdk 26 dan targetSdk 36. Maintainability dijaga melalui pemisahan layer core, data, domain, dan presentation.', first_line=True)
    add_table(doc, ['Ancaman', 'Risiko', 'Mitigasi pada SIRS'], [
        ['Unauthorized access', 'User membaca/mengubah laporan milik user lain.', 'Firestore Rules membatasi akses berdasarkan request.auth.uid dan role administrator.'],
        ['Privilege escalation', 'User biasa memakai fitur admin.', 'Route UI dipisahkan berdasarkan role dan Rules memvalidasi hak admin pada operasi sensitif.'],
        ['Malicious file upload', 'File bukan JPG/PNG/PDF disamarkan sebagai bukti.', 'Validasi MIME, ekstensi, ukuran, dan magic bytes pada FileValidation.kt.'],
        ['Secret leakage', 'API secret Cloudinary bocor dari APK.', 'Signature/delete dipindahkan ke Cloudflare Worker dengan Firebase JWT; Android hanya menyimpan cloud name/preset/folder.'],
        ['Data tampering', 'Status/severity laporan diubah sembarang.', 'Update status/severity hanya admin, perubahan dicatat pada statusHistories dan activityLogs.'],
        ['Abuse/spam', 'User membuat laporan/upload terlalu sering.', 'RateLimiter membatasi pembuatan laporan dan upload.']
    ])
    add_heading(doc, '2.2 Rancangan Arsitektur Sistem', 2)
    add_p(doc, 'Arsitektur SIRS menggunakan Android native sebagai client, Firebase sebagai backend utama, Cloud Firestore sebagai database, Firebase Authentication sebagai identitas, Firebase Messaging/Analytics/Crashlytics/App Check sebagai layanan pendukung, Cloudinary sebagai penyimpanan media, dan Cloudflare Worker sebagai backend serverless untuk operasi Cloudinary yang memerlukan signature.', first_line=True)
    code_block(doc, 'Gambar 1. Diagram Arsitektur Sistem\n\n[Android App: Jetpack Compose + ViewModel]\n        |\n        | Firebase Auth / Firestore SDK / FCM / Analytics\n        v\n[Firebase: Auth, Firestore, Messaging, Analytics, Crashlytics, App Check]\n        |\n        | Data laporan, user, kategori, log, status history\n        v\n[Cloud Firestore + Security Rules]\n\n[Android App] -- Firebase ID Token --> [Cloudflare Worker] -- signed params --> [Cloudinary]\n[Android App] -- upload file + signature -------------------------------> [Cloudinary]')
    code_block(doc, 'Gambar 2. DFD Level 0\n\nUser/Admin -> SIRS Android App -> Firebase Auth\nUser/Admin -> SIRS Android App -> Cloud Firestore\nSIRS Android App -> Cloudflare Worker -> Cloudinary\nAdmin -> SIRS Android App -> Status/Response/Kategori/Log')
    code_block(doc, 'Gambar 3. DFD Level 1 Pembuatan Laporan\n\n1. User mengisi form laporan.\n2. ViewModel memvalidasi field wajib dan rate limit.\n3. Jika ada attachment, FileValidation mengecek MIME, ekstensi, ukuran, dan magic bytes.\n4. Android meminta signature upload ke Worker memakai Firebase ID Token.\n5. File diupload ke Cloudinary.\n6. Repository membuat reportCode, SLA deadline, dan dokumen incidentReports.\n7. Activity log dibuat untuk audit.')
    add_heading(doc, '2.3 Rancangan Database', 2)
    add_p(doc, 'SIRS tidak menggunakan SQLite lokal, Room, SQLCipher, atau MySQL. Database yang dipakai adalah Cloud Firestore. Struktur data dibagi menjadi beberapa collection utama sebagai berikut.', first_line=True)
    add_table(doc, ['Collection/Subcollection', 'Isi Data', 'Keterangan Relasi'], [
        ['users', 'uid, nama, email, role, active, lastLoginAt, createdAt', 'Satu dokumen untuk setiap akun Firebase Authentication.'],
        ['categories', 'nama, slug, deskripsi, guidance, active', 'Dipakai form laporan dan dikelola admin.'],
        ['incidentReports', 'reportCode, userId, userName, title, category, severity, status, description, location, incidentDate, attachment, slaDeadline, adminResponse', 'Entitas utama laporan insiden.'],
        ['incidentReports/{reportId}/statusHistories', 'status, severity, response, changedBy, changedAt', 'Riwayat perubahan laporan oleh admin.'],
        ['activityLogs', 'action, actorId, actorName, targetId, metadata, createdAt', 'Audit aktivitas penting aplikasi.'],
        ['deviceTokens', 'userId, token, platform, updatedAt', 'Token FCM perangkat pengguna.']
    ])
    add_heading(doc, '2.4 Rancangan Antarmuka Pengguna (UI/UX)', 2)
    add_p(doc, 'Antarmuka dibuat menggunakan Jetpack Compose dan Material 3. Screen utama yang dirancang meliputi login, register, forgot password, dashboard user, dashboard admin, daftar laporan, detail laporan, form buat/edit laporan, daftar laporan admin, detail laporan admin, manajemen kategori, activity log, dan profil.', first_line=True)
    add_table(doc, ['Layar', 'Fungsi', 'Komponen UI Utama'], [
        ['Login/Register/Forgot Password', 'Autentikasi akun dan reset password.', 'TextField, Button, loading/error state, navigasi.'],
        ['User Dashboard', 'Ringkasan laporan milik pengguna dan shortcut fitur.', 'Card statistik, action card, top bar.'],
        ['Admin Dashboard', 'Ringkasan seluruh laporan dan akses modul admin.', 'Card statistik, action card, LazyColumn.'],
        ['Create/Edit Report', 'Input data laporan dan attachment.', 'OutlinedTextField, dropdown kategori/severity, date picker, attachment picker.'],
        ['Report List/Detail', 'Daftar, filter, detail, riwayat status, attachment, response.', 'LazyColumn, chip filter, card detail, badge status/severity.'],
        ['Category Management', 'CRUD kategori laporan.', 'Form kategori, switch active, list kategori.'],
        ['Activity Log', 'Menampilkan audit aktivitas.', 'LazyColumn dan activity log card.']
    ])
    doc.add_page_break()


def bab3(doc):
    add_heading(doc, 'BAB III\nIMPLEMENTASI DAN PEMBAHASAN', 1)
    add_heading(doc, '3.1 Lingkungan Pengembangan', 2)
    add_table(doc, ['Komponen', 'Spesifikasi'], [
        ['IDE', 'Android Studio'], ['Bahasa Pemrograman', 'Kotlin 2.2.10'], ['Minimum SDK', 'API 26'], ['Target SDK', 'API 36'], ['Compile SDK', 'API 36'], ['Build System', 'Gradle Kotlin DSL, Android Gradle Plugin 9.1.1'], ['Library Utama', 'Jetpack Compose, Material 3, Hilt, Firebase Auth, Firestore, Messaging, Analytics, Crashlytics, App Check, Coil, OkHttp, Coroutines'], ['Perangkat Pengujian', 'Emulator/perangkat Android API minimal 26']
    ])
    add_heading(doc, '3.2 Implementasi Antarmuka Pengguna (UI)', 2)
    add_p(doc, 'UI aplikasi berada pada folder app/src/main/java/com/adit/sirs/presentation. Theme utama berada pada presentation/theme/Theme.kt dan navigation graph berada pada presentation/navigation/AppNavGraph.kt. Setiap screen menerima state dari ViewModel, lalu mengirim event seperti login, submit laporan, filter, atau update status kembali ke ViewModel.', first_line=True)
    code_block(doc, 'Potongan alur Compose utama:\nMainActivity -> setContent { SIRSTheme { AppNavGraph() } }\nAppNavGraph -> memilih start destination berdasarkan currentUser dan role\nScreen Compose -> observe StateFlow dari ViewModel -> render Material 3 components')
    add_table(doc, ['Screen', 'File Implementasi', 'Pembahasan'], [
        ['Login', 'presentation/auth/LoginScreen.kt', 'Input email/password, tombol login, navigasi register dan forgot password.'],
        ['Register', 'presentation/auth/RegisterScreen.kt', 'Input nama, email, password, validasi registrasi.'],
        ['Dashboard User', 'presentation/dashboard/UserDashboardScreen.kt', 'Menampilkan ringkasan laporan user dan akses buat/lihat laporan.'],
        ['Dashboard Admin', 'presentation/dashboard/AdminDashboardScreen.kt', 'Menampilkan statistik seluruh laporan dan akses modul admin.'],
        ['Create Report', 'presentation/reports/CreateReportScreen.kt', 'Form laporan, dropdown kategori/severity, tanggal, lokasi, deskripsi, attachment.'],
        ['Report Detail', 'presentation/reports/ReportDetailScreen.kt', 'Menampilkan detail laporan, attachment, response admin, dan status history.'],
        ['Admin Report Detail', 'presentation/admin/AdminReportDetailScreen.kt', 'Admin mengubah status, severity, response, dan menghapus laporan.']
    ])
    add_heading(doc, '3.3 Implementasi Firebase Authentication dan Cloud Firestore', 2)
    add_p(doc, 'Firebase Authentication dipakai untuk login, registrasi, reset password, dan session pengguna. Implementasi data source berada pada FirebaseAuthDataSource.kt, repository pada AuthRepositoryImpl.kt, dan state UI pada AuthViewModel.kt. Setelah login, profil pengguna diambil dari collection users untuk menentukan role user atau administrator.', first_line=True)
    code_block(doc, 'Alur autentikasi:\nAuthViewModel.login(email, password)\n -> AuthRepositoryImpl.login(...)\n -> FirebaseAuthDataSource.signIn(...)\n -> FirestoreUserDataSource.getUser(uid)\n -> AppNavGraph mengarahkan role administrator ke admin_dashboard dan role user ke user_dashboard')
    add_p(doc, 'Cloud Firestore digunakan untuk menyimpan users, categories, incidentReports, statusHistories, activityLogs, dan deviceTokens. Akses data dilakukan melalui data source pada package data/remote dan dipetakan ke domain model melalui mapper pada package data/mapper.', first_line=True)
    add_table(doc, ['Fitur Firestore', 'File Utama', 'Operasi'], [
        ['User profile', 'FirestoreUserDataSource.kt', 'Membaca/menulis profil dan role pengguna.'],
        ['Laporan insiden', 'FirestoreReportDataSource.kt', 'Create, observe, update, delete laporan dan status history.'],
        ['Kategori', 'FirestoreCategoryDataSource.kt', 'Observe kategori aktif, CRUD kategori admin.'],
        ['Activity log', 'FirestoreActivityLogDataSource.kt', 'Create dan observe log aktivitas.'],
        ['FCM token', 'FcmTokenDataSource.kt', 'Menyimpan/menghapus token perangkat.']
    ])
    add_p(doc, 'Security Rules pada firestore.rules membatasi operasi berdasarkan autentikasi, kepemilikan dokumen, role administrator, status laporan, dan validitas data. User hanya dapat mengakses laporan miliknya sendiri, sedangkan admin dapat mengelola seluruh laporan dan kategori.', first_line=True)
    add_heading(doc, '3.4 Implementasi Cloudinary dan Cloudflare Worker', 2)
    add_p(doc, 'Upload bukti laporan menggunakan Cloudinary. Sebelum upload, aplikasi memvalidasi file melalui FileValidation.kt dan MimeTypeValidator.kt. Validasi mencakup MIME whitelist, kecocokan ekstensi, batas ukuran, dan magic bytes untuk JPG, PNG, serta PDF.', first_line=True)
    code_block(doc, 'Alur upload bukti:\nReportViewModel.setAttachment(uri)\n -> FileValidation.validateFile(contentResolver, uri)\n -> ReportViewModel.submitReport()\n -> CloudinaryUploadDataSource meminta signature ke Worker dengan Firebase ID Token\n -> Android upload file ke Cloudinary memakai signature\n -> URL/public_id disimpan pada field attachment di incidentReports')
    add_p(doc, 'Cloudinary API secret tidak disimpan di Android. BuildConfig Android hanya berisi cloud name, upload preset, folder upload, dan endpoint delete/signature. Operasi yang membutuhkan signature diproses oleh Cloudflare Worker yang memverifikasi Firebase JWT sebelum membuat signature.', first_line=True)
    add_heading(doc, '3.5 Fitur-Fitur Utama Aplikasi', 2)
    add_table(doc, ['No', 'Fitur', 'Implementasi Teknis'], [
        [1, 'Login dan registrasi', 'AuthViewModel, AuthRepositoryImpl, FirebaseAuthDataSource, FirestoreUserDataSource.'],
        [2, 'Reset password', 'Firebase Authentication reset email melalui ForgotPasswordScreen/AuthViewModel.'],
        [3, 'Dashboard user/admin', 'DashboardViewModel menghitung statistik laporan dan screen menampilkan card statistik.'],
        [4, 'Pembuatan laporan', 'ReportViewModel memvalidasi form, upload attachment, hitung SLA, lalu ReportRepositoryImpl menyimpan data.'],
        [5, 'Daftar/detail/filter laporan user', 'ObserveUserReports, status filter, severity filter, ReportListScreen, ReportDetailScreen.'],
        [6, 'Edit/hapus laporan user', 'Dibatasi untuk pemilik laporan dan status pending sesuai Rules.'],
        [7, 'Daftar/detail/filter laporan admin', 'AdminViewModel observeAllReports, search query, filter status/severity/tanggal.'],
        [8, 'Update status/severity/response', 'AdminViewModel.updateReportStatus menulis laporan, statusHistories, dan activityLogs.'],
        [9, 'Manajemen kategori', 'Admin CRUD/toggle kategori; user membaca kategori aktif.'],
        [10, 'Activity log', 'FirestoreActivityLogDataSource mencatat aktivitas penting dan ActivityLogScreen menampilkan log.'],
        [11, 'SLA deadline', 'SlaCalculator menentukan batas waktu penanganan berdasarkan severity.'],
        [12, 'FCM token', 'SirsFirebaseMessagingService dan FcmTokenDataSource menyiapkan token perangkat.']
    ])
    add_heading(doc, '3.6 Pengujian Aplikasi', 2)
    add_p(doc, 'Pengujian dilakukan melalui inspeksi source code, sinkronisasi dokumentasi dengan file project, dan build Gradle. Hasil build CLI yang diverifikasi adalah ./gradlew assembleDebug berhasil pada project SIRS.', first_line=True)
    add_table(doc, ['Skenario', 'Langkah Uji', 'Hasil Diharapkan', 'Hasil Aktual'], [
        ['Login user/admin', 'Masukkan email dan password valid.', 'User masuk ke dashboard sesuai role.', 'Sesuai desain AuthViewModel dan AppNavGraph.'],
        ['Registrasi user', 'Isi nama, email, password, submit.', 'Akun Auth dan profil Firestore dibuat.', 'Didukung AuthRepositoryImpl.'],
        ['Buat laporan', 'Isi form lengkap dan submit.', 'incidentReports dibuat dengan status pending.', 'Didukung ReportViewModel dan FirestoreReportDataSource.'],
        ['Upload bukti valid', 'Pilih JPG/PNG/PDF valid.', 'File lolos validasi dan diupload ke Cloudinary.', 'Didukung FileValidation dan CloudinaryUploadDataSource.'],
        ['Upload bukti tidak valid', 'Pilih file dengan MIME/ekstensi/signature tidak sesuai.', 'File ditolak sebelum upload.', 'Didukung FileValidation.'],
        ['Update status admin', 'Admin membuka detail dan mengubah status.', 'Status berubah dan status history tercatat.', 'Didukung AdminViewModel dan FirestoreReportDataSource.'],
        ['Akses data user lain', 'User mencoba membaca/mengubah laporan bukan miliknya.', 'Ditolak oleh Firestore Rules.', 'Rules membatasi owner/admin.'],
        ['Manajemen kategori', 'Admin tambah/edit/nonaktif kategori.', 'Kategori berubah dan user hanya melihat kategori aktif.', 'Didukung CategoryRepository dan Rules.']
    ])
    add_table(doc, ['Kontrol Keamanan', 'File Bukti', 'Penjelasan'], [
        ['Firebase Authentication', 'FirebaseAuthDataSource.kt, AuthViewModel.kt', 'Identitas pengguna dan session login.'],
        ['Role-based access', 'FirestoreUserDataSource.kt, AppNavGraph.kt, firestore.rules', 'Role administrator/user menentukan akses UI dan database.'],
        ['Firestore Security Rules', 'firestore.rules', 'Validasi akses collection users, categories, incidentReports, statusHistories, activityLogs, deviceTokens.'],
        ['Validasi upload', 'FileValidation.kt, MimeTypeValidator.kt', 'Mencegah file berbahaya dan spoofing tipe file.'],
        ['Secret isolation', 'CloudinaryUploadDataSource.kt, cloudinary-backend-worker/', 'Signature dibuat Worker, bukan hardcoded di APK.'],
        ['Audit trail', 'FirestoreActivityLogDataSource.kt', 'Aktivitas penting tercatat untuk admin.'],
        ['Rate limiting', 'RateLimiter.kt', 'Membatasi pembuatan laporan/upload agar tidak disalahgunakan.']
    ])
    doc.add_page_break()


def bab4(doc):
    add_heading(doc, 'BAB IV\nKESIMPULAN DAN SARAN', 1)
    add_heading(doc, '4.1 Kesimpulan', 2)
    add_numbers(doc, [
        'Project SIRS berhasil dirancang sebagai aplikasi Android native berbasis Kotlin dan Jetpack Compose untuk pelaporan serta respons insiden keamanan siber.',
        'Aplikasi menerapkan autentikasi pengguna dan pemisahan hak akses user/admin menggunakan Firebase Authentication, profil role pada Firestore, navigasi berbasis role, dan Firestore Security Rules.',
        'Struktur database Cloud Firestore mendukung kebutuhan utama aplikasi melalui collection users, categories, incidentReports, statusHistories, activityLogs, dan deviceTokens.',
        'Keamanan upload bukti diterapkan melalui validasi MIME, ekstensi, ukuran, magic bytes, serta signed upload Cloudinary melalui Cloudflare Worker sehingga API secret tidak berada di aplikasi Android.',
        'Fitur utama seperti dashboard, laporan, filter, detail, edit/hapus pending, update status, response admin, manajemen kategori, activity log, dan SLA deadline telah dipetakan ke file implementasi project.'
    ])
    add_heading(doc, '4.2 Saran', 2)
    add_bullets(doc, [
        'Menambahkan notifikasi FCM yang lebih lengkap ketika status laporan berubah atau response admin diberikan.',
        'Menambahkan pengujian instrumentasi otomatis untuk flow login, pembuatan laporan, dan update status admin.',
        'Menambahkan export laporan ke PDF untuk kebutuhan dokumentasi insiden.',
        'Menambahkan dashboard analitik insiden berdasarkan kategori, severity, dan tren waktu.',
        'Menambahkan opsi biometric authentication untuk aksi sensitif seperti delete laporan atau akses dashboard admin.'
    ])
    doc.add_page_break()


def refs_appendix(doc):
    add_heading(doc, 'DAFTAR PUSTAKA', 1)
    refs = [
        '[1] Android Developers, “Jetpack Compose Documentation,” developer.android.com.',
        '[2] Firebase Documentation, “Firebase Authentication, Cloud Firestore, Firebase Security Rules, Firebase App Check,” firebase.google.com/docs.',
        '[3] Cloudinary Documentation, “Image and file upload API,” cloudinary.com/documentation.',
        '[4] Cloudflare Developers, “Workers Documentation,” developers.cloudflare.com/workers.',
        '[5] OWASP, “Mobile Application Security Verification Standard (MASVS),” owasp.org.'
    ]
    for r in refs: add_p(doc, r, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_heading(doc, 'LAMPIRAN', 1)
    add_heading(doc, 'Lampiran 1 – Link Source Code', 2)
    add_p(doc, 'Repository/Google Drive source code: ........................................')
    add_heading(doc, 'Lampiran 2 – Screenshot Lengkap Aplikasi', 2)
    add_p(doc, 'Screenshot aplikasi dapat dilampirkan dari layar Login, Register, Dashboard User, Dashboard Admin, Create Report, Report List, Report Detail, Admin Report List, Admin Report Detail, Category Management, Activity Log, dan Profile.')
    add_heading(doc, 'Lampiran 3 – Source Code Penting', 2)
    add_bullets(doc, [
        'app/src/main/java/com/adit/sirs/presentation/navigation/AppNavGraph.kt',
        'app/src/main/java/com/adit/sirs/presentation/reports/ReportViewModel.kt',
        'app/src/main/java/com/adit/sirs/presentation/admin/AdminViewModel.kt',
        'app/src/main/java/com/adit/sirs/core/security/FileValidation.kt',
        'app/src/main/java/com/adit/sirs/data/remote/CloudinaryUploadDataSource.kt',
        'firestore.rules',
        'cloudinary-backend-worker/sirs-cloudinary/src/index.ts'
    ])
    add_heading(doc, 'Lampiran 4 – Dokumentasi API Backend Cloudinary Worker', 2)
    add_table(doc, ['Endpoint', 'Method', 'Fungsi'], [
        ['/generate-signature', 'POST', 'Membuat signature upload Cloudinary setelah memverifikasi Firebase ID Token.'],
        ['/delete', 'POST', 'Menghapus asset Cloudinary setelah memverifikasi Firebase ID Token dan hak akses.'],
        ['/health', 'GET', 'Health check Worker.']
    ])


def main():
    raise SystemExit(
        'Deprecated: jangan gunakan generator ini untuk laporan berbasis template DOCX. '
        'Script ini membangun ulang dokumen dan dapat merusak struktur/formatting template. '
        'Gunakan workflow docx-template-preservation: mulai dari Template Laporan Akhir Projek UAS.docx, '
        'preserve cover/TOC/heading/section, lalu replace placeholder atau insert content pada anchor.'
    )

if __name__ == '__main__':
    main()
