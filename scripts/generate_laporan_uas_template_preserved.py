from copy import deepcopy
from pathlib import Path
from zipfile import ZipFile
import re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / 'Template Laporan Akhir Projek UAS.docx'
OUTPUT = ROOT / 'Laporan Akhir Projek UAS - SIRS TEMPLATE PRESERVED v6.docx'
SCREENSHOTS = ROOT / 'docs' / 'screenshots'

PROJECT_TITLE = 'SIRS - Sistem Informasi Respons Insiden Keamanan Siber'
STUDENT_NAME = 'Nama Mahasiswa'
STUDENT_ID_CLASS = 'NPM/Kelas'

SCREENSHOT_ITEMS = [
    ('01_login_awal.png', 'Gambar 1. Layar login awal aplikasi SIRS.'),
    ('02_register_terisi.png', 'Gambar 2. Form registrasi user dengan validasi password.'),
    ('03_dashboard_user.png', 'Gambar 3. Dashboard pengguna setelah registrasi.'),
    ('04_form_buat_laporan_user.png', 'Gambar 4. Form pembuatan laporan insiden.'),
    ('05_daftar_laporan_user.png', 'Gambar 5. Daftar laporan pada role user.'),
    ('07_dashboard_admin.png', 'Gambar 6. Dashboard administrator.'),
    ('08_daftar_laporan_admin.png', 'Gambar 7. Daftar laporan pada role admin.'),
    ('09_manajemen_kategori.png', 'Gambar 8. Manajemen kategori laporan.'),
    ('10_log_aktivitas.png', 'Gambar 9. Activity log aplikasi.'),
]

SECURITY_SCREENSHOT_ITEMS = [
    ('security/security_01_register_weak_password_filled.png', 'Gambar 10. Form registrasi diisi dengan password lemah sebagai skenario uji validasi keamanan password.'),
    ('security/security_02_register_weak_password_rejected.png', 'Gambar 11. Aplikasi menolak registrasi karena password tidak memenuhi standar minimal keamanan.'),
    ('security/security_03_invalid_login_filled.png', 'Gambar 12. Form login diisi dengan kredensial tidak valid sebagai skenario uji autentikasi.'),
    ('security/security_04_invalid_login_rejected.png', 'Gambar 13. Firebase Authentication menolak kredensial login yang tidak valid.'),
    ('security/security_05_admin_login_filled.png', 'Gambar 14. Form login admin diisi untuk menguji autentikasi dan role administrator.'),
    ('security/security_06_admin_dashboard_role_verified.png', 'Gambar 15. Dashboard admin terbuka setelah role administrator tervalidasi.'),
    ('security/security_07_activity_log_admin_only.png', 'Gambar 16. Activity log admin sebagai bukti fitur audit trail dan pembatasan akses role admin.'),
]

SECTION_CONTENT = {
    'Latar Belakang': [
        'Ancaman keamanan siber terhadap pengguna aplikasi mobile semakin meningkat, terutama pada proses pelaporan insiden yang sering dilakukan secara tidak terstruktur melalui percakapan informal atau dokumen terpisah. Kondisi tersebut menyulitkan pencatatan kronologi, pemantauan status, penentuan prioritas, dan dokumentasi bukti pendukung.',
        'SIRS dikembangkan sebagai aplikasi Android native untuk membantu pengguna melaporkan insiden keamanan siber dan membantu administrator menindaklanjuti laporan secara terukur. Aplikasi menyediakan autentikasi, dashboard berbasis role, pembuatan laporan, upload bukti, pengelolaan status, kategori, SLA, dan audit log.',
        'Tema ini dipilih karena relevan dengan kebutuhan respons insiden: laporan harus tercatat, hanya dapat diakses oleh pihak berwenang, dan perubahan status harus dapat diaudit. Penerapan keamanan dilakukan pada autentikasi Firebase, Firestore Security Rules, validasi file, pembatasan rate, App Check, serta delegasi operasi sensitif Cloudinary melalui backend Cloudflare Worker.'
    ],
    'Rumusan Masalah': [
        'Bagaimana membangun aplikasi Android native untuk pelaporan dan penanganan insiden keamanan siber berbasis role user dan administrator?',
        'Bagaimana mengintegrasikan Firebase Authentication dan Cloud Firestore agar data laporan hanya dapat diakses oleh pemilik laporan dan administrator?',
        'Bagaimana menerapkan validasi keamanan upload bukti laporan dan pencatatan aktivitas penting aplikasi?',
        'Bagaimana menyediakan dashboard, filter, status penanganan, kategori, dan SLA agar proses respons insiden lebih terstruktur?'
    ],
    'Tujuan': [
        'Membangun aplikasi SIRS berbasis Kotlin dan Jetpack Compose untuk registrasi, login, pelaporan insiden, pemantauan status, serta administrasi laporan.',
        'Mengimplementasikan integrasi Firebase Authentication dan Cloud Firestore dengan pemisahan akses berdasarkan role user dan administrator.',
        'Menerapkan validasi file berbasis MIME, ekstensi, magic bytes, pembatasan rate, dan upload Cloudinary melalui signature yang didelegasikan ke Cloudflare Worker.',
        'Menyediakan fitur dashboard, filter, kategori, activity log, dan perhitungan batas waktu penanganan berdasarkan tingkat keparahan laporan.'
    ],
    'Manfaat': [
        'Bagi pengguna akhir, aplikasi memudahkan pelaporan insiden keamanan siber lengkap dengan kategori, severity, lokasi, deskripsi, tanggal kejadian, dan bukti pendukung.',
        'Bagi administrator, aplikasi membantu memantau seluruh laporan, memfilter laporan, memperbarui status/severity, memberi respons, mengelola kategori, dan membaca audit log.',
        'Bagi mahasiswa sebagai developer, proyek ini melatih penerapan Android native modern, MVVM/Clean Architecture, Firebase, validasi keamanan file, dan dokumentasi teknis aplikasi.'
    ],
    'Batasan Masalah': [
        'Aplikasi berjalan pada Android minimum SDK 26 dengan target SDK 36 dan dikembangkan menggunakan Kotlin serta Jetpack Compose.',
        'Penyimpanan utama menggunakan Cloud Firestore, bukan SQLite lokal/Room Database. Karena itu pembahasan SQLCipher dijelaskan sebagai tidak diterapkan pada proyek ini.',
        'Backend utama menggunakan Firebase dan Cloudflare Worker untuk delegasi operasi Cloudinary; proyek tidak menggunakan server MySQL/PHP.',
        'Fokus keamanan proyek berada pada autentikasi Firebase, role access, keamanan data Firestore, validasi upload, Cloudflare Worker, dan audit log.'
    ],
    'Analisis Kebutuhan Sistem': [
        'Kebutuhan fungsional aplikasi meliputi registrasi dan login, reset password, dashboard user/admin, pembuatan laporan insiden, upload bukti, daftar/detail laporan, filter status dan severity, perubahan status oleh admin, respons admin, riwayat status, manajemen kategori, activity log, penyimpanan token FCM, dan perhitungan SLA.',
        'Kebutuhan data mencakup profil pengguna, role administrator/user, kategori laporan, data laporan insiden, metadata attachment, riwayat perubahan status, log aktivitas, dan token perangkat. Setiap data dipisahkan ke collection Firestore agar alur baca/tulis dapat dikontrol melalui Security Rules dan repository aplikasi.',
        'Kebutuhan non-fungsional meliputi kompatibilitas Android API 26 ke atas, UI responsif dengan Jetpack Compose Material 3, pemisahan layer agar kode mudah dirawat, validasi input, pembatasan rate untuk aksi sensitif, observasi data realtime, dan pengamanan akses data melalui Firebase Authentication serta Firestore Security Rules.',
        'Analisis ancaman mencakup akses tidak sah antar pengguna, manipulasi status laporan oleh non-admin, upload file berbahaya, penyalahgunaan endpoint eksternal, kebocoran secret Cloudinary, penghapusan lampiran tanpa otorisasi, dan kurangnya jejak audit. Mitigasi dilakukan melalui role-based rules, validasi file, Cloudflare Worker, Firebase ID Token, App Check, rate limiter, dan activity log.'
    ],
    'Rancangan Arsitektur Sistem': [
        'Arsitektur SIRS menggunakan pola MVVM/Clean Architecture ringan. Layer presentation berisi screen Jetpack Compose dan ViewModel, layer domain berisi model serta kontrak repository, layer data berisi implementasi repository, data source Firebase/Cloudinary, DTO, dan mapper, sedangkan layer core berisi validator, helper, konstanta, dan utilitas keamanan.',
        'Alur data utama adalah Screen Compose -> ViewModel -> Repository Interface -> Repository Implementation -> Remote Data Source -> Firebase/Cloudinary -> Mapper -> Domain Model -> StateFlow -> UI. Firebase Authentication menangani identitas pengguna, Cloud Firestore menyimpan users, categories, incidentReports, statusHistories, activityLogs, dan deviceTokens.',
        'Operasi upload bukti menggunakan validasi lokal terlebih dahulu, kemudian aplikasi meminta signature ke Cloudflare Worker dengan Firebase ID Token. Worker memverifikasi token sebelum membuat signature Cloudinary sehingga API Secret tidak pernah disimpan di aplikasi Android.'
    ],
    'Rancangan Database (Jika ada)': [
        'Project ini tidak menggunakan SQLite lokal, Room Database, SQLCipher, maupun MySQL. Database utama yang digunakan adalah Cloud Firestore karena aplikasi membutuhkan sinkronisasi laporan antara pengguna dan administrator tanpa server database mandiri.',
        'Struktur Firestore terdiri dari collection users untuk profil dan role, categories untuk kategori laporan, incidentReports untuk data laporan utama, subcollection statusHistories untuk riwayat perubahan status per laporan, activityLogs untuk audit aktivitas penting, dan deviceTokens untuk token FCM perangkat.',
        'Field penting pada laporan meliputi reportCode, userId, userName, title, categoryId/categoryName, severity, status, description, location, incidentDate, attachment, adminResponse, slaDeadline, handledBy, handledAt, createdAt, updatedAt, dan deletedAt. Metadata attachment menyimpan originalName, publicId, secureUrl, resourceType, format, mimeType, bytes, dan uploadedAt.',
        'Aturan akses database diatur melalui firestore.rules. User hanya dapat membaca detail laporan miliknya dan mengubah/menghapus laporan saat status masih pending, sedangkan administrator dapat membaca seluruh laporan, mengubah status/severity/respons, mengelola kategori, serta membaca activity log. Fallback rules menolak akses collection lain yang tidak didefinisikan.'
    ],
    'Rancangan Antarmuka Pengguna (UI/UX)': [
        'Rancangan UI menggunakan Jetpack Compose dan Material 3 dengan pembagian layar utama: login, register, forgot password, dashboard user, dashboard admin, daftar laporan, form laporan, detail laporan, manajemen kategori, activity log, dan profil.',
        'Komponen utama yang dipakai meliputi TextField, Button, Card, LazyColumn, Scaffold, TopAppBar, FilterChip, dialog/bottom sheet, badge status, badge severity, serta komponen reusable untuk loading, empty state, dan error view.',
        'UI dibedakan berdasarkan role. User diarahkan ke dashboard pengguna dan fitur laporan miliknya, sedangkan administrator diarahkan ke dashboard admin, daftar seluruh laporan, kategori, dan activity log.'
    ],
    'Lingkungan Pengembangan': [
        'Lingkungan pengembangan SIRS menggunakan Android Studio, Gradle Kotlin DSL, Kotlin 2.2.10, Android Gradle Plugin 9.1.1, Jetpack Compose Material 3, Java 17, compile SDK 36, minimum SDK 26, Firebase, Hilt, KSP, Coroutines, dan Cloudinary/Cloudflare Worker untuk upload bukti.'
    ],
    'Implementasi Antarmuka Pengguna (UI)': [
        'Antarmuka pengguna diimplementasikan dengan Jetpack Compose dan Material 3. LoginScreen, RegisterScreen, dan ForgotPasswordScreen menangani autentikasi. UserDashboardScreen menampilkan ringkasan laporan milik pengguna. ReportListScreen, CreateReportScreen, EditReportScreen, dan ReportDetailScreen menangani siklus laporan pada sisi user.',
        'Pada sisi administrator, AdminDashboardScreen menampilkan statistik global, AdminReportListScreen menyediakan daftar/pencarian/filter laporan, AdminReportDetailScreen menyediakan perubahan status, severity, dan respons, CategoryManagementScreen mengelola kategori laporan, sedangkan ActivityLogScreen menampilkan jejak aktivitas sistem.',
        'Setiap screen membaca state dari ViewModel melalui StateFlow. Perubahan input form diteruskan ke ViewModel, sedangkan hasil operasi Firebase/Cloudinary direpresentasikan sebagai state loading, success, atau error sehingga UI dapat memberi umpan balik kepada pengguna tanpa mengakses data source secara langsung.'
    ],
    'Implementasi Cloud Firestore dan Security Rules': [
        'Project ini memetakan kebutuhan database aplikasi ke Cloud Firestore, bukan SQLite lokal. Firestore dipilih karena SIRS membutuhkan sinkronisasi data laporan antara user dan administrator secara realtime, termasuk status laporan, riwayat perubahan, kategori, device token, dan activity log.',
        'Struktur database terdiri dari collection users, categories, incidentReports, activityLogs, deviceTokens, serta subcollection statusHistories pada setiap laporan. Collection users menyimpan profil dan role, incidentReports menyimpan data laporan insiden, statusHistories menyimpan riwayat perubahan, categories menyimpan kategori, deviceTokens menyimpan token FCM, dan activityLogs menyimpan jejak audit.',
        'Security Rules menggunakan helper signedIn, isOwner, currentUserDoc, isAdmin, isActiveUser, validUserCreate, validReportCreate, ownsExistingReport, onlyPendingOwnerEdit, dan canDeleteExistingReport. Validasi create laporan memastikan userId sesuai UID login, status awal pending, deletedAt null, panjang title/deskripsi/lokasi sesuai batas, dan severity hanya low/medium/high/critical.',
        'Hak akses dibatasi sesuai role: user hanya dapat membaca detail laporan miliknya serta mengubah/menghapus laporan pending, sedangkan administrator dapat membaca seluruh laporan, mengubah status/severity/respons, mengelola kategori, dan membaca activity log. Dengan rancangan ini, fungsi database terenkripsi lokal pada template digantikan oleh kontrol akses server-side yang sesuai dengan arsitektur aktual project.'
    ],
    'Implementasi Backend Cloudflare Worker dan Cloudinary Secure Upload': [
        'Project ini tidak menggunakan server MySQL/PHP. Kebutuhan koneksi backend eksternal dipenuhi melalui Cloudflare Worker untuk operasi sensitif Cloudinary, sedangkan data utama tetap berada di Cloud Firestore.',
        'Saat user mengunggah bukti laporan, aplikasi memvalidasi MIME, ekstensi, ukuran maksimal 2 MB, dan magic bytes file terlebih dahulu melalui FileValidation. Setelah itu CloudinaryUploadDataSource mengambil Firebase ID Token dari user aktif dan meminta signature upload ke endpoint Cloudflare Worker /generate-signature.',
        'Worker memverifikasi token sebelum membuat signature Cloudinary. Android hanya menerima api_key, timestamp, folder, dan signature, lalu melakukan multipart upload ke Cloudinary. API Secret tetap berada di sisi Worker dan tidak disimpan pada BuildConfig, gradle.properties, maupun source code Android.',
        'Penghapusan attachment juga didelegasikan ke endpoint backend dengan header Authorization: Bearer Firebase ID Token. Endpoint menerima publicId dan resourceType, kemudian Worker memakai secret server-side untuk operasi destroy Cloudinary. Pola ini menggantikan bagian integrasi MySQL pada template dengan secure backend delegation yang benar-benar digunakan project.'
    ],
    'Implementasi Firebase Cloud Firestore + Security Rules': [
        'Project ini menggunakan Cloud Firestore, bukan Firebase Realtime Database. Integrasi Firebase mencakup Firebase Authentication, Cloud Firestore, Firebase Messaging, Analytics, Crashlytics, dan App Check.',
        'Firebase Authentication dipakai untuk login, registrasi, reset password, dan identitas pengguna. Setelah registrasi, profil user disimpan di collection users dengan role user. Admin dibedakan melalui role administrator pada profil Firestore.',
        'Cloud Firestore menyimpan data users, categories, incidentReports, statusHistories, activityLogs, dan deviceTokens. Security Rules membatasi akses berdasarkan status login, UID pemilik laporan, status laporan, dan role administrator sehingga user tidak dapat membaca laporan milik user lain.'
    ],

    'Fitur-Fitur Utama Aplikasi': [
        'Fitur utama SIRS terdiri dari login, registrasi, reset password, dashboard user/admin, pembuatan laporan, upload bukti, daftar laporan, detail laporan, edit/hapus laporan pending oleh pemilik, filter status/severity, pencarian admin, update status/severity, respons admin, hapus laporan sesuai rules, riwayat status, manajemen kategori, panduan kategori, activity log, penyimpanan token FCM, dan SLA deadline.',
        'Pada sisi user, alur utama dimulai dari registrasi/login, membuka dashboard, membuat laporan insiden, melihat daftar laporan, membuka detail, mengedit atau menghapus laporan jika masih pending, lalu memantau respons admin. Pada sisi admin, alur utama meliputi memantau dashboard, mencari/memfilter laporan, membuka detail, mengubah status/severity, memberi respons, mengelola kategori, dan membaca log aktivitas.',
        'Fitur yang bersifat pendukung namun tetap relevan adalah analytics event dasar, Crashlytics dependency, App Check, Firestore offline persistence, dan callback FCM. Beberapa komponen tersebut sudah disiapkan di kode sebagai fondasi operasional, sedangkan notifikasi push kustom masih menjadi ruang pengembangan lanjutan.'
    ],
    'Pengujian Aplikasi': [
        'Pengujian fungsional dilakukan dalam dua tahap. Tahap pertama adalah build verification menggunakan ./gradlew assembleDebug untuk memastikan source code, dependency, resource, dan konfigurasi Gradle dapat dikompilasi. Tahap kedua adalah pengujian runtime pada emulator Android emulator-5554 dengan aplikasi debug yang diinstal dan dijalankan langsung.',
        'Skenario user yang diuji meliputi tampilan login, registrasi akun, login user, dashboard pengguna, pembuatan laporan, pengisian judul/kategori/severity/deskripsi/lokasi/tanggal, daftar laporan, dan akses detail laporan. Skenario administrator yang diuji meliputi login admin, dashboard admin, daftar seluruh laporan, akses detail laporan, manajemen kategori, dan pembacaan activity log.',
        'Pengujian keamanan didokumentasikan sebagai pemeriksaan source-code dan runtime terbatas terhadap kontrol yang benar-benar ada di project: Firebase Authentication, pemisahan role user/administrator, Firestore Security Rules, validasi input laporan, validasi upload file berbasis MIME/ekstensi/ukuran/magic bytes, rate limiter, App Check, penggunaan Firebase ID Token untuk Cloudflare Worker, dan tidak adanya API Secret Cloudinary pada source Android.',
        'Batasan pengujian juga dicatat agar redaksi tetap jujur: dependency test otomatis JUnit/Espresso/Compose sudah tersedia, namun test case otomatis khusus belum menjadi bukti utama. Karena itu bukti utama laporan adalah hasil build, inspeksi source, screenshot emulator, dan verifikasi manual alur aplikasi. Pengujian lanjutan yang disarankan adalah Firebase Emulator Suite untuk rules dan instrumented UI test untuk alur kritikal.'
    ],
    'Kesimpulan': [
        'SIRS berhasil dibangun sebagai aplikasi Android native untuk pelaporan dan penanganan insiden keamanan siber dengan pemisahan role user dan administrator.',
        'Integrasi Firebase Authentication dan Cloud Firestore berhasil digunakan untuk autentikasi, penyimpanan laporan, kategori, riwayat status, dan audit log dengan pembatasan akses melalui Security Rules.',
        'Validasi upload bukti dan delegasi signature Cloudinary melalui Cloudflare Worker memperkuat keamanan karena file diperiksa sebelum upload dan secret eksternal tidak disimpan di aplikasi Android.',
        'Dashboard, filter, status penanganan, kategori, activity log, dan SLA membantu proses respons insiden menjadi lebih terstruktur dan terdokumentasi.'
    ],
    'Saran': [
        'Menambahkan notifikasi push yang lebih lengkap agar user mendapat pemberitahuan saat status laporan berubah.',
        'Menambahkan export laporan ke PDF/CSV dan dashboard analitik insiden untuk kebutuhan evaluasi keamanan.',
        'Melakukan pengujian Security Rules lebih luas dengan Firebase Emulator Suite dan menambahkan automated UI test untuk alur kritikal.'
    ],
    'Daftar Pustaka (Jika ada)': [
        '[1] Google, “Android Developers Documentation,” https://developer.android.com/.',
        '[2] Google, “Firebase Documentation,” https://firebase.google.com/docs.',
        '[3] Google, “Cloud Firestore Security Rules,” https://firebase.google.com/docs/firestore/security/get-started.',
        '[4] Cloudinary, “Image and Video Upload API,” https://cloudinary.com/documentation/image_upload_api_reference.',
        '[5] OWASP, “Mobile Application Security,” https://owasp.org/www-project-mobile-top-10/.'
    ],
    'Lampiran 1 – Link Source Code: Tautan Google Drive yang berisi seluruh source code proyek/ Repository GitHub': [
        'Source code berada pada folder project lokal: C:\\Users\\Code X\\AndroidStudioProjects\\SIRS. Apabila dikumpulkan melalui Google Drive atau GitHub, tautan repository dapat ditempelkan pada bagian ini.'
    ],
    'Lampiran 3 – Source Code: Listing kode penting aplikasi SIRS': [
        'Bagian source code penting meliputi AppNavGraph.kt untuk navigasi role, AuthViewModel.kt untuk autentikasi, ReportViewModel.kt untuk alur laporan, AdminViewModel.kt untuk alur admin, FileValidation.kt untuk validasi bukti, CloudinaryUploadDataSource.kt untuk upload, dan firestore.rules untuk kontrol akses database.'
    ],
}

BULLET_SECTIONS = {'Rumusan Masalah', 'Tujuan', 'Manfaat', 'Batasan Masalah', 'Kesimpulan', 'Saran'}


def copy_run_format(src, dst):
    if src is None:
        return
    dst.bold = src.bold
    dst.italic = src.italic
    dst.underline = src.underline
    dst.font.name = src.font.name
    dst.font.size = src.font.size
    if src._element.rPr is not None:
        dst._element.get_or_add_rPr().rPr = deepcopy(src._element.rPr)


def replace_paragraph_text(paragraph, text):
    src_run = next((r for r in paragraph.runs if r.text), paragraph.runs[0] if paragraph.runs else None)
    for r in list(paragraph.runs):
        r.text = ''
    run = paragraph.add_run(text)
    copy_run_format(src_run, run)


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def insert_paragraph_after(paragraph, text='', style=None):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    new_para._element = new_p
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_table_after(paragraph, rows, cols):
    tmp_doc = paragraph._parent
    try:
        table = tmp_doc.add_table(rows=rows, cols=cols)
    except TypeError:
        table = tmp_doc.add_table(rows=rows, cols=cols, width=Inches(6.0))
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    paragraph._p.addnext(tbl)
    return table


def clone_body_format(target, sample):
    if sample is not None:
        ppr = sample.find(qn('w:pPr'))
        if ppr is not None:
            current = target._p.find(qn('w:pPr'))
            if current is not None:
                target._p.remove(current)
            target._p.insert(0, deepcopy(ppr))
    else:
        target.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        target.paragraph_format.line_spacing = 1.5
        target.paragraph_format.space_after = Pt(6)


def style_inserted(paragraph, sample_para=None, bullet=False):
    clone_body_format(paragraph, sample_para)
    if bullet:
        # Preserve native Word list style already present in template where possible.
        paragraph.style = 'List Paragraph'
    if not paragraph.runs:
        paragraph.add_run('')
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(12)


def set_run_times_new_roman(run):
    run.font.name = 'Times New Roman'
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ('ascii', 'hAnsi', 'eastAsia', 'cs'):
        rfonts.set(qn(f'w:{attr}'), 'Times New Roman')


def normalize_all_fonts(doc):
    for style in doc.styles:
        if style.type in (WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER, WD_STYLE_TYPE.TABLE):
            try:
                style.font.name = 'Times New Roman'
                rpr = style.element.get_or_add_rPr()
                rfonts = rpr.get_or_add_rFonts()
                for attr in ('ascii', 'hAnsi', 'eastAsia', 'cs'):
                    rfonts.set(qn(f'w:{attr}'), 'Times New Roman')
            except Exception:
                pass
    containers = [doc]
    for section in doc.sections:
        containers.extend([section.header, section.footer, section.first_page_header, section.first_page_footer])
    for container in containers:
        for p in container.paragraphs:
            for run in p.runs:
                set_run_times_new_roman(run)
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            set_run_times_new_roman(run)


def norm(text):
    return re.sub(r'\s+', ' ', text or '').strip()


def find_paragraph(doc, text):
    target = norm(text)
    for p in doc.paragraphs:
        if norm(p.text) == target:
            return p
    raise ValueError(f'Paragraph not found: {text}')


def following_until_next_heading(doc, idx):
    out = []
    for p in doc.paragraphs[idx + 1:]:
        if p.style.name.startswith('Heading'):
            break
        out.append(p)
    return out


def paragraph_index(doc, paragraph):
    for i, p in enumerate(doc.paragraphs):
        if p._p is paragraph._p:
            return i
    raise ValueError(f'Paragraph not in document: {paragraph.text!r}')


def replace_section(doc, heading_text, items):
    heading_aliases = {
        'Implementasi Cloud Firestore dan Security Rules': 'Implementasi SQLite + SQLCipher (Jika ada)',
        'Implementasi Backend Cloudflare Worker dan Cloudinary Secure Upload': 'Implementasi Integrasi MySQL (Secure Connection) (Jika ada)',
        'Implementasi Firebase Cloud Firestore + Security Rules': 'Implementasi Firebase Realtime Database + Security Rules (Jika ada)',
        'Lampiran 3 – Source Code: Listing kode penting aplikasi SIRS': 'Lampiran 3 – Source Code: Listing kode lengkap untuk bagian-bagian yang paling penting (misalnya: database class, security rules, biometric manager class, dll.).',
    }
    original_heading_text = heading_aliases.get(heading_text, heading_text)
    heading = find_paragraph(doc, original_heading_text)
    if original_heading_text != heading_text:
        replace_paragraph_text(heading, heading_text)
    idx = paragraph_index(doc, heading)
    old = following_until_next_heading(doc, idx)
    sample = next((deepcopy(p._p) for p in old if p.text.strip() and not p.style.name.startswith('List')), None)
    for p in old:
        if p.text.strip():
            delete_paragraph(p)
    anchor = heading
    for item in reversed(items):
        p = insert_paragraph_after(anchor, item)
        style_inserted(p, sample, bullet=heading_text in BULLET_SECTIONS)
    return find_paragraph(doc, heading_text)


def remove_section(doc, heading_text):
    heading = find_paragraph(doc, heading_text)
    idx = paragraph_index(doc, heading)
    body = following_until_next_heading(doc, idx)
    for p in body:
        delete_paragraph(p)
    delete_paragraph(heading)


def fill_environment_table(doc):
    if not doc.tables:
        return
    table = doc.tables[0]
    data = [
        ('IDE', 'Android Studio'),
        ('Bahasa Pemrograman', 'Kotlin 2.2.10'),
        ('Minimum SDK', 'API 26 (Android 8.0 Oreo)'),
        ('Target SDK', 'API 36'),
        ('Database Lokal', 'Tidak menggunakan SQLite/Room; data utama menggunakan Cloud Firestore'),
        ('Layanan Tambahan', 'Firebase Auth, Cloud Firestore, FCM, Analytics, Crashlytics, App Check, Cloudinary, Cloudflare Worker'),
    ]
    for i, (a, b) in enumerate(data, start=1):
        if i >= len(table.rows):
            table.add_row()
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b


def add_testing_table_after(doc, heading_text):
    heading = find_paragraph(doc, heading_text)
    # Insert after all content currently under heading.
    idx = paragraph_index(doc, heading)
    body = following_until_next_heading(doc, idx)
    anchor = body[-1] if body else heading
    cap = insert_paragraph_after(anchor, 'Tabel 2. Hasil Pengujian Aplikasi SIRS', style='Caption')
    table = insert_table_after(cap, 1, 4)
    try:
        table.style = doc.tables[0].style
    except Exception:
        pass
    headers = ['Skenario', 'Langkah Uji', 'Hasil Diharapkan', 'Hasil Aktual']
    for j,h in enumerate(headers):
        table.rows[0].cells[j].text = h
    rows = [
        ('Registrasi user', 'Isi nama, email, password valid, lalu tekan Daftar', 'Akun user dibuat dan diarahkan ke dashboard', 'Berhasil pada emulator'),
        ('Login admin', 'Masuk dengan akun administrator', 'Admin masuk ke dashboard admin', 'Berhasil pada emulator'),
        ('Buat laporan', 'User mengisi form laporan insiden', 'Laporan tersimpan dengan status pending', 'Berhasil diuji melalui UI'),
        ('Manajemen kategori', 'Admin membuka menu kategori', 'Kategori tampil dan dapat dikelola', 'Berhasil ditampilkan'),
        ('Activity log', 'Admin membuka log aktivitas', 'Aktivitas penting tercatat', 'Berhasil ditampilkan'),
    ]
    for row in rows:
        cells = table.add_row().cells
        for j,val in enumerate(row):
            cells[j].text = val


def add_security_testing_table_after(doc, heading_text):
    heading = find_paragraph(doc, heading_text)
    idx = paragraph_index(doc, heading)
    body = following_until_next_heading(doc, idx)
    anchor = body[-1] if body else heading
    cap = insert_paragraph_after(anchor, 'Tabel 3. Dokumentasi Pengujian Keamanan Aplikasi SIRS', style='Caption')
    table = insert_table_after(cap, 1, 5)
    try:
        table.style = doc.tables[0].style
    except Exception:
        pass
    headers = ['Aspek Keamanan', 'Metode Pengujian', 'Bukti Implementasi', 'Hasil yang Diharapkan', 'Status']
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
    rows = [
        ('Autentikasi Firebase', 'Inspeksi AuthViewModel/FirebaseAuthDataSource dan uji login pada emulator', 'AuthViewModel.kt, FirebaseAuthDataSource.kt, screenshot login/user/admin', 'Hanya akun valid yang dapat masuk; role menentukan dashboard', 'Terdokumentasi dan berhasil diuji'),
        ('Pemisahan role user/admin', 'Inspeksi AppNavGraph, UserRole, users collection, dan Firestore rules isAdmin', 'AppNavGraph.kt, UserRole.kt, firestore.rules', 'User biasa tidak diarahkan ke fitur admin; admin dapat membuka fitur pengelolaan', 'Terdokumentasi'),
        ('Kontrol akses Firestore', 'Audit firestore.rules untuk users, categories, incidentReports, statusHistories, activityLogs, dan deviceTokens', 'firestore.rules helper signedIn/isAdmin/validReportCreate/onlyPendingOwnerEdit', 'User hanya akses data miliknya; admin mengelola data sesuai role; collection lain default deny', 'Terdokumentasi'),
        ('Validasi input laporan', 'Inspeksi validReportCreate dan konstanta batas panjang title/deskripsi/lokasi/severity', 'firestore.rules, AppConstants.kt, ReportFormState.kt', 'Data laporan tidak valid ditolak sebelum atau saat masuk database', 'Terdokumentasi'),
        ('Validasi upload file', 'Inspeksi FileValidation untuk MIME, ekstensi, ukuran 2 MB, dan magic bytes JPG/PNG/PDF', 'FileValidation.kt, MimeTypeValidator.kt, AppConstants.kt', 'File palsu/berbahaya dan ukuran berlebih ditolak sebelum upload', 'Terdokumentasi'),
        ('Cloudinary secret protection', 'Inspeksi BuildConfig dan CloudinaryUploadDataSource; pastikan signature/delete memakai Firebase ID Token ke Worker', 'CloudinaryUploadDataSource.kt, app/build.gradle.kts', 'API Secret tidak berada di aplikasi Android; operasi sensitif terjadi di backend', 'Terdokumentasi'),
        ('Rate limiting aksi sensitif', 'Inspeksi RateLimiter dan pemanggilan pada ReportViewModel', 'RateLimiter.kt, ReportViewModel.kt', 'Spam pembuatan laporan/upload dibatasi di sisi aplikasi', 'Terdokumentasi'),
        ('Audit log', 'Inspeksi FirestoreActivityLogDataSource dan ActivityLogScreen; uji tampilan log admin', 'FirestoreActivityLogDataSource.kt, ActivityLogScreen.kt, screenshot activity log', 'Aktivitas penting tercatat dan hanya admin yang membaca log', 'Berhasil ditampilkan'),
    ]
    for row in rows:
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = val


def add_screenshots_under(doc, heading_text):
    heading = find_paragraph(doc, heading_text)
    # Remove only placeholder body under screenshot appendix if any.
    idx = paragraph_index(doc, heading)
    body = following_until_next_heading(doc, idx)
    for p in body:
        if p.text.strip():
            delete_paragraph(p)
    anchor = heading
    note = insert_paragraph_after(anchor, 'Screenshot berikut diambil langsung dari emulator Android bawaan melalui ADB saat aplikasi debug SIRS berjalan. Gambar disisipkan pada lampiran tanpa mengubah struktur cover, daftar isi, BAB, dan heading asli template.')
    style_inserted(note, None)
    anchor = note
    for filename, caption in SCREENSHOT_ITEMS:
        img = SCREENSHOTS / filename
        if not img.exists():
            raise FileNotFoundError(img)
        p_img = insert_paragraph_after(anchor, '')
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(str(img), width=Inches(3.0))
        cap = insert_paragraph_after(p_img, caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.style = 'Caption'
        for r in cap.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(11)
            r.italic = True
        anchor = cap

    sec = insert_paragraph_after(anchor, 'Screenshot Pengujian Keamanan Aplikasi')
    style_inserted(sec, None)
    for r in sec.runs:
        r.bold = True
    anchor = sec
    note_security = insert_paragraph_after(anchor, 'Screenshot berikut merupakan bukti runtime dari emulator untuk skenario pengujian keamanan yang dapat diverifikasi langsung melalui UI, yaitu validasi password lemah, penolakan kredensial login tidak valid, verifikasi dashboard admin berdasarkan role, dan akses activity log sebagai bukti audit trail. Pengujian keamanan yang bersifat server-side seperti Firestore Security Rules, validasi magic bytes, dan delegasi Cloudinary tetap dijelaskan melalui tabel serta inspeksi source karena tidak seluruhnya terlihat sebagai layar UI.')
    style_inserted(note_security, None)
    anchor = note_security
    for filename, caption in SECURITY_SCREENSHOT_ITEMS:
        img = SCREENSHOTS / filename
        if not img.exists():
            raise FileNotFoundError(img)
        p_img = insert_paragraph_after(anchor, '')
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(str(img), width=Inches(3.0))
        cap = insert_paragraph_after(p_img, caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.style = 'Caption'
        for r in cap.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(11)
            r.italic = True
        anchor = cap


def add_toc_update_setting(doc):
    settings = doc.settings.element
    update = settings.find(qn('w:updateFields'))
    if update is None:
        update = OxmlElement('w:updateFields')
        settings.append(update)
    update.set(qn('w:val'), 'true')


def main():
    if not TEMPLATE.exists():
        raise FileNotFoundError(TEMPLATE)
    doc = Document(str(TEMPLATE))

    # Cover placeholder replacement only; cover layout paragraphs remain untouched.
    replacements = {
        '“Judul / Nama Aplikasi yang dibuat”': PROJECT_TITLE,
        'Nama Mahasiswa': STUDENT_NAME,
        'NPM/Kelas': STUDENT_ID_CLASS,
    }
    for p in doc.paragraphs:
        if p.text.strip() in replacements:
            replace_paragraph_text(p, replacements[p.text.strip()])

    for heading, items in SECTION_CONTENT.items():
        replace_section(doc, heading, items)

    # Remove optional template-only sections that have no actual SIRS implementation
    # and no genuinely related project scope.
    remove_section(doc, 'Implementasi Biometric Authentication (Jika ada)')
    remove_section(doc, 'Lampiran 4 – Dokumentasi REST API (jika menggunakan MySQL): Daftar endpoint API, method HTTP, parameter, dan contoh response')

    fill_environment_table(doc)
    # Insert security table first because both testing tables anchor to the same
    # last paragraph; later insertions appear closer to the anchor in Word XML.
    add_security_testing_table_after(doc, 'Pengujian Aplikasi')
    add_testing_table_after(doc, 'Pengujian Aplikasi')
    add_screenshots_under(doc, 'Lampiran 2 – Screenshot Lengkap Aplikasi: Kumpulan screenshot seluruh layar aplikasi yang belum sempat dicantumkan dalam bab sebelumnya')
    add_toc_update_setting(doc)
    normalize_all_fonts(doc)
    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == '__main__':
    main()
