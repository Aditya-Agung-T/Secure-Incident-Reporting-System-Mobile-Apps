from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
INPUT = ROOT / 'Laporan Akhir Projek UAS - SIRS.docx'
OUTPUT = ROOT / 'Laporan Akhir Projek UAS - SIRS dengan Screenshot.docx'
SS = ROOT / 'docs' / 'screenshots'

shots = [
    ('01_login_awal.png', 'Gambar 4. Layar login awal aplikasi SIRS.'),
    ('02_register_terisi.png', 'Gambar 5. Form registrasi user uji terisi lengkap dan password memenuhi standar keamanan.'),
    ('03_dashboard_user.png', 'Gambar 6. Dashboard pengguna setelah registrasi akun user berhasil.'),
    ('04_form_buat_laporan_user.png', 'Gambar 7. Form pembuatan laporan insiden pada role user.'),
    ('05_daftar_laporan_user.png', 'Gambar 8. Daftar laporan milik pengguna.'),
    ('06_login_admin_terisi.png', 'Gambar 9. Form login admin terisi dengan kredensial role administrator.'),
    ('07_dashboard_admin.png', 'Gambar 10. Dashboard admin setelah login berhasil.'),
    ('08_daftar_laporan_admin.png', 'Gambar 11. Daftar seluruh laporan pada role admin.'),
    ('09_manajemen_kategori.png', 'Gambar 12. Layar manajemen kategori laporan oleh admin.'),
    ('10_log_aktivitas.png', 'Gambar 13. Layar activity log sebagai bukti audit aktivitas aplikasi.'),
]

def fmt(p, align=None, italic=False, bold=False, size=12):
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    if align:
        p.alignment = align
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        r.font.size = Pt(size)
        r.italic = italic
        r.bold = bold

def add_heading(doc, text):
    p=doc.add_paragraph()
    r=p.add_run(text)
    r.bold=True
    fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT)

def add_note(doc, text):
    p=doc.add_paragraph()
    p.add_run(text)
    fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

def add_shot(doc, filename, caption):
    path=SS/filename
    if not path.exists():
        raise FileNotFoundError(path)
    p=doc.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(3.0))
    cap=doc.add_paragraph()
    cap.add_run(caption)
    fmt(cap, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=11)


def main():
    doc=Document(str(INPUT))
    add_heading(doc, 'Lampiran 2 – Screenshot Lengkap Aplikasi dari Emulator')
    add_note(doc, 'Screenshot berikut diambil langsung dari emulator Android bawaan melalui ADB setelah aplikasi debug SIRS di-install dan diuji. State layar diverifikasi menggunakan UI Automator dump sebelum gambar dimasukkan ke laporan. Akun user dibuat melalui flow registrasi aplikasi, sedangkan akun admin menggunakan kredensial yang disediakan pada pengujian.')
    for fn, cap in shots:
        add_shot(doc, fn, cap)
    doc.save(str(OUTPUT))
    print(OUTPUT)

if __name__ == '__main__':
    main()
